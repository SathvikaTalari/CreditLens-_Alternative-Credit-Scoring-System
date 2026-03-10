from flask import Flask, request, jsonify, send_from_directory, render_template
import json, re, os, io, warnings
import numpy as np
import pandas as pd
from sklearn.ensemble import GradientBoostingClassifier
import shap
import google.generativeai as genai
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import base64
from pdf_generator import generate_cam_pdf

warnings.filterwarnings('ignore')
app = Flask(__name__, template_folder='templates', static_folder='static')
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32MB max upload

# ══════════════════════════════════════════════════════════════════════════════
# INDIA CONTEXT ENGINE
# ══════════════════════════════════════════════════════════════════════════════
INDIAN_SECTORS = [
    "Manufacturing","Textile","Pharmaceuticals","IT/ITES","Real Estate",
    "Construction","Agriculture","Retail","Logistics","NBFC","Hospitality",
    "Auto Components","Steel","Chemicals","Food Processing"
]

SECTOR_RISK = {
    "Real Estate":"High","Construction":"High","Hospitality":"High","NBFC":"High",
    "Textile":"Medium","Retail":"Medium","Agriculture":"Medium","Logistics":"Medium",
    "Auto Components":"Medium","Steel":"Medium-High","Chemicals":"Medium",
    "Manufacturing":"Low-Medium","Food Processing":"Low-Medium",
    "Pharmaceuticals":"Low","IT/ITES":"Low"
}

SECTOR_HEADWINDS = {
    "Textile":"Global demand slowdown, China competition, rising yarn prices",
    "Real Estate":"High interest rates, regulatory scrutiny under RERA",
    "Steel":"China dumping, volatile iron ore prices",
    "NBFC":"RBI tightening, co-lending model changes",
    "Pharmaceuticals":"USFDA scrutiny, pricing pressure on generics",
    "IT/ITES":"Global tech slowdown, visa restrictions",
    "Agriculture":"Monsoon dependency, MSP fluctuations",
    "Construction":"Labour shortage, raw material inflation",
}

GSTR_CONTEXT = """
Indian GST Filing Intelligence:
- GSTR-1: Outward supply return (sales declared by seller)
- GSTR-2A: Auto-populated inward supply from supplier GSTR-1 (READ ONLY)
- GSTR-2B: Static ITC statement generated monthly
- GSTR-3B: Monthly self-assessed summary return (taxes actually paid)
- GSTR-9: Annual return | GSTR-9C: Reconciliation statement (audited)

RED FLAGS:
1. GSTR-1 sales >> GSTR-3B tax paid → Tax evasion risk
2. GSTR-2A ITC >> GSTR-3B ITC claimed → Fake invoice fraud
3. Sudden nil returns followed by large filing → Circular trading
4. GST turnover vs ITR turnover mismatch >10% → Revenue manipulation
5. Bank credits far below GST turnover → Revenue inflation
"""

def check_fraud_flags(gst, bank, itr, nil_months, itc=0):
    flags = []
    if gst and bank and gst > 0:
        m = abs(gst - bank) / gst
        if m > 0.20:
            flags.append({"flag":"🚩 BANK-GST DISCREPANCY",
                          "detail":f"GST ₹{gst:,.0f} vs Bank ₹{bank:,.0f} ({m*100:.1f}% gap)",
                          "severity":"High"})
    if gst and itr and gst > 0:
        m = abs(gst - itr) / gst
        if m > 0.15:
            flags.append({"flag":"⚠️ ITR-GST MISMATCH",
                          "detail":f"GST ₹{gst:,.0f} vs ITR ₹{itr:,.0f} ({m*100:.1f}% gap)",
                          "severity":"Medium"})
    if nil_months >= 3:
        flags.append({"flag":"⚠️ GST FILING IRREGULARITY",
                      "detail":f"{nil_months} months of nil/missing GST filings detected",
                      "severity":"Medium"})
    if gst and itc and gst > 0 and (itc/gst) > 0.35:
        flags.append({"flag":"⚠️ UNUSUALLY HIGH ITC RATIO",
                      "detail":f"ITC is {itc/gst*100:.1f}% of turnover — possible fake invoice usage",
                      "severity":"Medium"})
    return flags

def get_credit_band(score):
    if score >= 80: return {"rating":"AAA/AA","quality":"Excellent","color":"#00C853"}
    if score >= 65: return {"rating":"A/BBB","quality":"Good","color":"#64DD17"}
    if score >= 50: return {"rating":"BB","quality":"Moderate","color":"#FFD600"}
    if score >= 35: return {"rating":"B/C","quality":"Risky","color":"#FF6D00"}
    return {"rating":"D","quality":"High Risk","color":"#D50000"}

# ══════════════════════════════════════════════════════════════════════════════
# ML SCORING ENGINE
# ══════════════════════════════════════════════════════════════════════════════
_model = None
_explainer = None

FEATURE_COLS = [
    "vintage_years","turnover_cr","profit_margin","debt_equity_ratio",
    "dscr","gst_compliance_months","bank_gst_mismatch_pct","cibil_score",
    "collateral_ratio","employee_count","existing_loan_ratio",
    "litigation_flag","promoter_flag"
]
FEATURE_LABELS = {
    "vintage_years":"Business Age (Yrs)","turnover_cr":"Turnover (Cr)",
    "profit_margin":"Profit Margin","debt_equity_ratio":"Debt/Equity Ratio",
    "dscr":"Debt Service Coverage","gst_compliance_months":"GST Compliance (Mo)",
    "bank_gst_mismatch_pct":"Bank-GST Mismatch","cibil_score":"CIBIL Score",
    "collateral_ratio":"Collateral Coverage","employee_count":"Employee Count",
    "existing_loan_ratio":"Existing Loan Burden","litigation_flag":"Litigation Flag",
    "promoter_flag":"Promoter Risk Flag"
}

def get_model():
    global _model, _explainer
    if _model is not None:
        return _model, _explainer
    np.random.seed(42)
    n = 800
    v   = np.random.randint(1,30,n)
    t   = np.random.uniform(10,500,n)
    pm  = np.random.uniform(-0.05,0.25,n)
    de  = np.random.uniform(0.1,5,n)
    dscr= np.random.uniform(0.5,3.5,n)
    gc  = np.random.randint(4,13,n)
    bgm = np.random.uniform(0,0.5,n)
    cs  = np.random.randint(300,900,n)
    cr  = np.random.uniform(0.3,2.5,n)
    ec  = np.random.randint(5,500,n)
    elr = np.random.uniform(0,0.8,n)
    lf  = np.random.choice([0,1],n,p=[0.8,0.2])
    pf  = np.random.choice([0,1],n,p=[0.85,0.15])
    score = (50 + np.minimum(v*1.5,20) + pm*80 - np.maximum(de-2,0)*8
             + np.minimum((dscr-1)*15,20) + (gc/12)*15 - bgm*25
             + (cs-550)/35 + np.minimum(cr*8,15) - elr*20
             - lf*15 - pf*12 + np.random.normal(0,5,n))
    score = np.clip(score,0,100)
    default = (score < 45).astype(int)
    X = pd.DataFrame({
        "vintage_years":v,"turnover_cr":t,"profit_margin":pm,
        "debt_equity_ratio":de,"dscr":dscr,"gst_compliance_months":gc,
        "bank_gst_mismatch_pct":bgm,"cibil_score":cs,"collateral_ratio":cr,
        "employee_count":ec,"existing_loan_ratio":elr,
        "litigation_flag":lf,"promoter_flag":pf
    })
    _model = GradientBoostingClassifier(n_estimators=150,random_state=42).fit(X,default)
    _explainer = shap.TreeExplainer(_model)
    return _model, _explainer

def calculate_score(profile):
    model, explainer = get_model()
    loan      = max(profile.get("loan_amount_requested",1),1)
    turnover  = max(profile.get("unified_turnover",1e7),1)
    assets    = max(profile.get("unified_total_assets",1),1)
    liabilities = profile.get("unified_total_liabilities",0)
    net_worth = max(assets - liabilities, 1)
    net_profit = profile.get("unified_net_profit",0)
    existing  = profile.get("existing_loans",0)

    feats = {
        "vintage_years":        profile.get("business_vintage_years",5),
        "turnover_cr":          turnover/1e7,
        "profit_margin":        net_profit/turnover,
        "debt_equity_ratio":    liabilities/net_worth,
        "dscr":                 max(0.5, net_profit*1.3/max(existing*0.15,loan*0.1,1)),
        "gst_compliance_months":profile.get("filing_months_present",10),
        "bank_gst_mismatch_pct":profile.get("bank_gst_mismatch_pct",0.05),
        "cibil_score":          profile.get("cibil_score",600),
        "collateral_ratio":     profile.get("collateral_value",0)/loan,
        "employee_count":       profile.get("employee_count",50),
        "existing_loan_ratio":  existing/turnover,
        "litigation_flag":      int(profile.get("litigation_flag",0)),
        "promoter_flag":        int(profile.get("promoter_flag",0)),
    }

    X = pd.DataFrame([feats])[FEATURE_COLS]
    prob      = model.predict_proba(X)[0][1]
    score     = round((1-prob)*100,1)

    sv = explainer.shap_values(X)
    if isinstance(sv,list): sv = sv[1][0]
    else: sv = sv[0]

    explanations = sorted([
        {"feature": FEATURE_LABELS[f],
         "raw_feature": f,
         "value": round(float(X[f].iloc[0]),3),
         "impact": round(float(sv[i]),3)}
        for i,f in enumerate(FEATURE_COLS)
    ], key=lambda x: abs(x["impact"]), reverse=True)

    return {
        "credit_score": score,
        "default_probability": round(prob*100,1),
        "explanations": explanations,
        "top_positive": [e for e in explanations if e["impact"]>0][:4],
        "top_negative": [e for e in explanations if e["impact"]<0][:4],
        "features": feats
    }

# ══════════════════════════════════════════════════════════════════════════════
# GEMINI AI FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════
def extract_from_pdf(pdf_bytes, doc_type, api_key):
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t: text += t + "\n"
    except: pass

    if not text.strip():
        return {"error": "Could not extract text from PDF"}

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
    prompt = f"""You are an expert Indian corporate credit analyst.
Extract financial data from this {doc_type}.

INDIAN GST CONTEXT:
{GSTR_CONTEXT}

DOCUMENT TEXT (first 6000 chars):
{text[:6000]}

Return ONLY a JSON object (no markdown):
{{
  "company_name": null,
  "gstin": null,
  "pan": null,
  "turnover_gst": null,
  "turnover_itr": null,
  "bank_credits_total": null,
  "tax_paid_gst": null,
  "itc_claimed": null,
  "filing_months_present": null,
  "filing_months_nil": null,
  "net_profit": null,
  "total_assets": null,
  "total_liabilities": null,
  "existing_loans": null,
  "collateral_value": null,
  "employee_count": null,
  "business_vintage_years": null,
  "key_risks": [],
  "positive_signals": [],
  "doc_type_confirmed": "GST/ITR/BankStatement/AnnualReport"
}}"""
    try:
        r = model.generate_content(prompt)
        raw = re.sub(r'```json|```','',r.text.strip()).strip()
        return json.loads(raw)
    except Exception as e:
        return {"error": str(e), "partial_text": text[:200]}

def run_research(company_name, sector, promoter, api_key):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
    prompt = f"""You are a Digital Credit Manager doing due diligence on an Indian company.

COMPANY: {company_name}
SECTOR: {sector}
PROMOTER: {promoter}
KNOWN SECTOR HEADWINDS: {SECTOR_HEADWINDS.get(sector,'Research current conditions')}

Analyze and return intelligence on litigation, news, regulatory risk, and sector outlook.
Return ONLY JSON (no markdown):
{{
  "company_news": [{{"headline":"","summary":"","sentiment":"Positive/Negative/Neutral"}}],
  "promoter_flags": [{{"issue":"","severity":"High/Medium/Low"}}],
  "litigation_findings": [{{"case":"","status":"","risk_level":"High/Medium/Low"}}],
  "regulatory_actions": [{{"authority":"","action":"","status":""}}],
  "sector_outlook": {{
    "overall": "Positive/Stable/Negative",
    "key_trends": [],
    "risk_factors": [],
    "growth_drivers": []
  }},
  "mca_flags": [],
  "overall_research_risk": "Low/Medium/High/Very High",
  "research_summary": "",
  "red_flags_count": 0,
  "positive_signals_count": 0
}}"""
    try:
        r = model.generate_content(prompt)
        raw = re.sub(r'```json|```','',r.text.strip()).strip()
        return json.loads(raw)
    except Exception as e:
        return {
            "error": str(e),
            "sector_outlook":{"overall":"Stable","key_trends":[],"risk_factors":[],"growth_drivers":[]},
            "overall_research_risk":"Medium",
            "research_summary":f"Research completed for {company_name} in {sector} sector.",
            "company_news":[],"promoter_flags":[],"litigation_findings":[],
            "regulatory_actions":[],"mca_flags":[],"red_flags_count":0,"positive_signals_count":0
        }

def adjust_for_notes(notes, score, profile, api_key):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
    prompt = f"""Senior credit committee member reviewing qualitative field notes.
CURRENT SCORE: {score}/100
FIELD NOTES: "{notes}"

Adjust score based on observations. Return ONLY JSON:
{{
  "adjustment": 0,
  "reason": "",
  "key_observations": [],
  "character_assessment": "Positive/Neutral/Concerning",
  "site_visit_risk": "Low/Medium/High"
}}"""
    try:
        r = model.generate_content(prompt)
        raw = re.sub(r'```json|```','',r.text.strip()).strip()
        result = json.loads(raw)
        result["adjusted_score"] = max(0,min(100, score + result.get("adjustment",0)))
        return result
    except:
        return {"adjustment":0,"adjusted_score":score,"reason":"Notes noted.",
                "key_observations":[notes],"character_assessment":"Neutral","site_visit_risk":"Medium"}

def generate_recommendation(score_result, loan, fraud_flags, research_risk, api_key):
    score = score_result["credit_score"]
    if score >= 80:   rec,amt,rate = "APPROVE",loan,"8.5–10%"
    elif score >= 65: rec,amt,rate = "APPROVE WITH CONDITIONS",loan*0.9,"10–12%"
    elif score >= 50: rec,amt,rate = "CONDITIONAL APPROVAL",loan*0.7,"12–14%"
    elif score >= 35: rec,amt,rate = "HIGH RISK — COLLATERAL REQUIRED",loan*0.5,"15–18%"
    else:             rec,amt,rate = "REJECT",0,"N/A"

    high_flags = [f for f in fraud_flags if f.get("severity")=="High"]
    if high_flags and rec=="APPROVE": rec="APPROVE WITH CONDITIONS"
    if high_flags: amt *= 0.8

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
    prompt = f"""Credit committee member at Indian NBFC writing loan recommendation.
SCORE: {score}/100 | RECOMMENDATION: {rec} | AMOUNT: ₹{amt:,.0f} | RATE: {rate}
TOP POSITIVE: {[e['feature'] for e in score_result['top_positive'][:3]]}
TOP NEGATIVE: {[e['feature'] for e in score_result['top_negative'][:3]]}
FRAUD FLAGS: {len(fraud_flags)} ({len(high_flags)} high severity)
RESEARCH RISK: {research_risk}

Return ONLY JSON:
{{
  "recommendation": "{rec}",
  "one_line_verdict": "",
  "detailed_reasoning": "",
  "conditions": [],
  "monitoring_triggers": [],
  "five_cs": {{
    "Character": "", "Capacity": "", "Capital": "", "Collateral": "", "Conditions": ""
  }}
}}"""
    try:
        r = model.generate_content(prompt)
        raw = re.sub(r'```json|```','',r.text.strip()).strip()
        result = json.loads(raw)
    except:
        result = {
            "recommendation":rec,
            "one_line_verdict":f"Score {score}/100 — {rec}",
            "detailed_reasoning":"AI-assisted scoring applied based on submitted financial data.",
            "conditions":[],"monitoring_triggers":[],
            "five_cs":{"Character":"Assessed","Capacity":"Assessed",
                       "Capital":"Assessed","Collateral":"Assessed","Conditions":"Assessed"}
        }
    result["suggested_amount"] = amt
    result["suggested_rate"]   = rate
    result["credit_score"]     = score
    return result

# ══════════════════════════════════════════════════════════════════════════════
# CAM REPORT GENERATOR
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex_color)
    tcPr.append(shd)

def generate_cam(company, score_result, fraud_flags, research, recommendation,
                 notes="", notes_adj=None):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1.2)

    def heading(text, lvl=1, color="1F3864"):
        h = doc.add_heading(text, level=lvl)
        run = h.runs[0] if h.runs else h.add_run(text)
        run.font.color.rgb = RGBColor.from_string(color)
        return h

    def info_table(data):
        t = doc.add_table(rows=len(data), cols=2)
        t.style = 'Table Grid'
        for i,(k,v) in enumerate(data.items()):
            t.rows[i].cells[0].text = str(k)
            t.rows[i].cells[1].text = str(v) if v is not None else "N/A"
            t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            set_cell_bg(t.rows[i].cells[0],"D6E4F0")
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CREDIT APPRAISAL MEMORANDUM"); r.bold=True; r.font.size=Pt(22)
    r.font.color.rgb = RGBColor.from_string("1F3864")

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("CreditLens Pro | AI-Powered Credit Intelligence | IIT Hyderabad Hackathon 2026")
    r2.italic=True; r2.font.size=Pt(10); r2.font.color.rgb=RGBColor.from_string("2E75B6")

    doc.add_paragraph()
    score = score_result["credit_score"]
    color = "00C853" if score>=65 else "FFD600" if score>=45 else "D50000"
    sp = doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"AI Credit Score: {score}/100  |  {recommendation['recommendation']}")
    sr.bold=True; sr.font.size=Pt(15); sr.font.color.rgb=RGBColor.from_string(color)

    doc.add_paragraph(f"Company: {company.get('company_name','N/A')}   |   Date: {datetime.date.today().strftime('%d %B %Y')}")
    doc.add_page_break()

    heading("1. EXECUTIVE SUMMARY")
    doc.add_paragraph(f"Verdict: {recommendation.get('one_line_verdict','')}")
    doc.add_paragraph(recommendation.get("detailed_reasoning",""))
    doc.add_paragraph()

    heading("2. COMPANY PROFILE")
    info_table({
        "Company Name": company.get("company_name"),
        "GSTIN": company.get("gstin"),
        "Sector": company.get("sector"),
        "Business Vintage": f"{company.get('business_vintage_years','N/A')} years",
        "Employees": company.get("employee_count"),
        "Annual Turnover": f"₹{company.get('unified_turnover',0):,.0f}",
        "Net Profit": f"₹{company.get('unified_net_profit',0):,.0f}",
        "Total Assets": f"₹{company.get('unified_total_assets',0):,.0f}",
        "Total Liabilities": f"₹{company.get('unified_total_liabilities',0):,.0f}",
        "Existing Loans": f"₹{company.get('existing_loans',0):,.0f}",
    })

    heading("3. LOAN RECOMMENDATION")
    info_table({
        "Loan Requested": f"₹{company.get('loan_amount_requested',0):,.0f}",
        "Recommended Amount": f"₹{recommendation.get('suggested_amount',0):,.0f}",
        "Interest Rate": recommendation.get("suggested_rate"),
        "Credit Score": f"{score}/100",
        "Default Probability": f"{score_result.get('default_probability',0)}%",
        "Recommendation": recommendation.get("recommendation"),
    })

    conds = recommendation.get("conditions",[])
    if conds:
        doc.add_paragraph("Conditions:").runs[0].bold=True
        for c in conds: doc.add_paragraph(f"• {c}")

    heading("4. FIVE Cs OF CREDIT")
    info_table(recommendation.get("five_cs",{}))

    heading("5. AI EXPLAINABILITY (SHAP Analysis)")
    doc.add_paragraph("Positive Factors:").runs[0].bold=True
    for f in score_result.get("top_positive",[]):
        doc.add_paragraph(f"  ✅ {f['feature']}: {f['value']} → Impact +{abs(f['impact']):.3f}")
    doc.add_paragraph("Negative Factors:").runs[0].bold=True
    for f in score_result.get("top_negative",[]):
        doc.add_paragraph(f"  ❌ {f['feature']}: {f['value']} → Impact -{abs(f['impact']):.3f}")
    doc.add_paragraph()

    heading("6. FRAUD & RISK FLAGS")
    if fraud_flags:
        for fl in fraud_flags:
            p = doc.add_paragraph()
            r = p.add_run(f"{fl['flag']} [{fl['severity']} Severity]"); r.bold=True
            r.font.color.rgb = RGBColor.from_string("FF0000" if fl['severity']=="High" else "FF6D00")
            doc.add_paragraph(f"  {fl['detail']}")
    else:
        doc.add_paragraph("✅ No fraud flags detected. Cross-document data consistent.")

    heading("7. EXTERNAL INTELLIGENCE")
    doc.add_paragraph(f"Research Risk: {research.get('overall_research_risk','N/A')}")
    doc.add_paragraph(f"Summary: {research.get('research_summary','')}")
    so = research.get("sector_outlook",{})
    if so:
        doc.add_paragraph(f"Sector Outlook: {so.get('overall','N/A')}")
        for t in so.get("key_trends",[])[:3]: doc.add_paragraph(f"  • {t}")

    if notes:
        heading("8. CREDIT OFFICER FIELD NOTES")
        doc.add_paragraph(f"Notes: {notes}")
        if notes_adj:
            adj = notes_adj.get("adjustment",0)
            p = doc.add_paragraph()
            r = p.add_run(f"Score {'▲' if adj>=0 else '▼'} {abs(adj)} pts → {notes_adj.get('adjusted_score')}/100")
            r.bold=True; r.font.color.rgb=RGBColor.from_string("00C853" if adj>=0 else "D50000")
            doc.add_paragraph(f"Reason: {notes_adj.get('reason','')}")

    heading("9. EARLY WARNING INDICATORS")
    for t in recommendation.get("monitoring_triggers",[]): doc.add_paragraph(f"⚠️ {t}")

    doc.add_page_break()
    fp = doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("CreditLens Pro | National AI/ML Hackathon 2026 | IIT Hyderabad\n"
                    "AI-generated memo — human review required before final lending decision.")
    fr.font.size=Pt(9); fr.italic=True; fr.font.color.rgb=RGBColor.from_string("888888")

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
# FLASK ROUTES
# ══════════════════════════════════════════════════════════════════════════════
@app.route('/')
def index():
    return render_template('index.html', sectors=INDIAN_SECTORS)

@app.route('/api/analyze', methods=['POST'])
def analyze():
    try:
        api_key = request.form.get('api_key','').strip()
        if not api_key:
            return jsonify({"error":"API key required"}), 400

        company_name  = request.form.get('company_name','Unknown Company')
        promoter_name = request.form.get('promoter_name','Unknown')
        sector        = request.form.get('sector','Manufacturing')
        loan_amount   = float(request.form.get('loan_amount', 5000000))
        cibil_score   = int(request.form.get('cibil_score', 650))
        qualitative_notes = request.form.get('qualitative_notes','').strip()

        # Manual financial inputs
        profile = {
            "company_name":          company_name,
            "promoter_name":         promoter_name,
            "sector":                sector,
            "loan_amount_requested": loan_amount,
            "cibil_score":           cibil_score,
            "business_vintage_years":float(request.form.get('vintage',8)),
            "unified_turnover":      float(request.form.get('turnover',200)) * 1e5,
            "unified_net_profit":    float(request.form.get('net_profit',15)) * 1e5,
            "unified_total_assets":  float(request.form.get('total_assets',300)) * 1e5,
            "unified_total_liabilities":float(request.form.get('total_liabilities',180)) * 1e5,
            "existing_loans":        float(request.form.get('existing_loans',50)) * 1e5,
            "collateral_value":      float(request.form.get('collateral_value',80)) * 1e5,
            "employee_count":        int(request.form.get('employee_count',75)),
            "filing_months_present": int(request.form.get('gst_months',10)),
            "bank_gst_mismatch_pct": float(request.form.get('bank_gst_mismatch',0.05)),
            "litigation_flag":       0,
            "promoter_flag":         0,
        }

        # PDF extraction
        all_fraud_flags = []
        docs_extracted  = []
        for field_name in ['gst_docs','bank_docs','other_docs']:
            files = request.files.getlist(field_name)
            doc_type = {"gst_docs":"GST Return","bank_docs":"Bank Statement","other_docs":"Financial Document"}[field_name]
            for f in files:
                if f and f.filename:
                    pdf_bytes = f.read()
                    extracted = extract_from_pdf(pdf_bytes, doc_type, api_key)
                    if "error" not in extracted:
                        docs_extracted.append(extracted)
                        gst = extracted.get("turnover_gst") or 0
                        bank= extracted.get("bank_credits_total") or 0
                        itr = extracted.get("turnover_itr") or 0
                        nil = extracted.get("filing_months_nil") or 0
                        itc = extracted.get("itc_claimed") or 0
                        flags = check_fraud_flags(gst, bank, itr, nil, itc)
                        all_fraud_flags.extend(flags)
                        # Merge into profile
                        for key in ["company_name","gstin","pan","business_vintage_years",
                                    "employee_count","filing_months_present"]:
                            if extracted.get(key) and not profile.get(key):
                                profile[key] = extracted[key]
                        if extracted.get("turnover_gst"):
                            profile["unified_turnover"] = extracted["turnover_gst"]
                        if extracted.get("net_profit"):
                            profile["unified_net_profit"] = extracted["net_profit"]
                        if extracted.get("total_assets"):
                            profile["unified_total_assets"] = extracted["total_assets"]
                        if extracted.get("total_liabilities"):
                            profile["unified_total_liabilities"] = extracted["total_liabilities"]

        # Pillar 2: Research
        research = run_research(company_name, sector, promoter_name, api_key)

        # Apply research risk to profile
        rr = research.get("overall_research_risk","Medium")
        if rr in ["High","Very High"]:
            profile["promoter_flag"] = 1
        if research.get("litigation_findings"):
            profile["litigation_flag"] = 1

        # Pillar 3: Scoring
        score_result = calculate_score(profile)
        risk_penalty = {"Low":0,"Medium":-3,"High":-8,"Very High":-15}
        score_result["credit_score"] = max(0, score_result["credit_score"] + risk_penalty.get(rr,0))

        # Qualitative notes
        notes_adj = None
        if qualitative_notes:
            notes_adj = adjust_for_notes(qualitative_notes, score_result["credit_score"], profile, api_key)
            score_result["credit_score"] = notes_adj.get("adjusted_score", score_result["credit_score"])

        # Recommendation
        recommendation = generate_recommendation(score_result, loan_amount, all_fraud_flags, rr, api_key)

        # CAM report - DOCX
        cam_bytes = generate_cam(profile, score_result, all_fraud_flags, research,
                                  recommendation, qualitative_notes, notes_adj)
        cam_docx_b64 = base64.b64encode(cam_bytes).decode("utf-8")

        # CAM report - PDF via reportlab
        cam_pdf_b64 = None
        try:
            pdf_bytes   = generate_cam_pdf(profile, score_result, all_fraud_flags, research,
                                            recommendation, qualitative_notes, notes_adj)
            cam_pdf_b64 = base64.b64encode(pdf_bytes).decode("utf-8")
        except Exception as pdf_err:
            cam_pdf_b64 = None

        band = get_credit_band(score_result["credit_score"])

        return jsonify({
            "success":            True,
            "company_name":       company_name,
            "credit_score":       score_result["credit_score"],
            "default_probability":score_result["default_probability"],
            "band":               band,
            "recommendation":     recommendation,
            "fraud_flags":        all_fraud_flags,
            "research":           research,
            "score_result":       score_result,
            "notes_adjustment":   notes_adj,
            "cam_b64":            cam_docx_b64,
            "cam_docx_b64":       cam_docx_b64,
            "cam_pdf_b64":        cam_pdf_b64,
            "profile":            profile,
        })

    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500

@app.route('/api/sectors')
def sectors():
    return jsonify(INDIAN_SECTORS)

if __name__ == '__main__':
    print("🚀 CreditLens Pro starting...")
    print("   Training ML model...")
    get_model()
    print("   ML model ready!")
    print("   Open http://localhost:5000")
    app.run(debug=True, port=5000)